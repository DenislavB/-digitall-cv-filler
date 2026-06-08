"""
CV Parser — extracts structured data from PDF or DOCX CVs.
Works completely offline (no API calls).
"""
import re
import os


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(filepath):
    import pdfplumber
    pages = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                pages.append(t)
    return "\n".join(pages)


def extract_text_from_docx(filepath):
    from docx import Document
    doc = Document(filepath)
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t and t not in lines:
                        lines.append(t)
    return "\n".join(lines)


def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        raw = extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        raw = extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

    if not raw or not raw.strip():
        raise RuntimeError(
            "No text could be extracted from this file.\n\n"
            "If it is a scanned/image PDF, the app cannot read it automatically.\n"
            "Please copy-paste the CV text manually into the form fields."
        )
    return raw


# ---------------------------------------------------------------------------
# Section detection  (very permissive — catches most real-world CV layouts)
# ---------------------------------------------------------------------------

# Each entry: (section_key, list_of_keyword_fragments)
# A line matches if it contains ANY of these fragments (case-insensitive)
# AND the line is short (< 60 chars), suggesting it's a heading.
SECTION_MAP = [
    ("summary",        ["summary", "profile", "about me", "objective", "overview",
                        "professional summary", "career objective", "personal statement"]),
    ("experience",     ["work experience", "experience", "employment history",
                        "professional experience", "career history", "work history",
                        "positions held", "career summary"]),
    ("education",      ["education", "academic", "qualifications", "academic background",
                        "studies", "academic history"]),
    ("skills",         ["skill", "technical skill", "competenc", "technolog",
                        "core competenc", "key skill", "expertise"]),
    ("languages",      ["language"]),
    ("certifications", ["certif", "course", "training", "license", "award",
                        "accreditation", "qualification"]),
]


def _is_section_header(line):
    """Return section key if this line looks like a section heading, else None."""
    stripped = line.strip()
    # Must be reasonably short and not look like a sentence
    if not stripped or len(stripped) > 70:
        return None
    # Reject lines that are clearly content (contain commas, bullets, long sentences)
    if stripped.count(",") > 2:
        return None
    low = stripped.lower()
    for key, keywords in SECTION_MAP:
        for kw in keywords:
            if kw in low:
                return key
    return None


def split_sections(raw_text):
    """Split raw CV text into named sections."""
    lines = [l.rstrip() for l in raw_text.splitlines()]
    sections = {"header": []}
    current = "header"

    for line in lines:
        key = _is_section_header(line)
        if key:
            current = key
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(line)

    return sections


# ---------------------------------------------------------------------------
# Field parsers
# ---------------------------------------------------------------------------

_CONTACT_RE = re.compile(r"[@+/|\\]|\d{5,}|linkedin|github|http|www\.", re.I)
_ADDRESS_RE = re.compile(
    r"(?i)^(home|address|street|road|avenue|ul\.|blvd|al\.|sq\.|floor|apt|"
    r"nationality|date of birth|gender|phone|email|mobile|tel:|fax|zip|postal|city:|"
    r"contacts?|profile picture)",
)
# Symbol/icon glyphs that PDF readers extract as unicode private-use characters
_ICON_RE = re.compile(r"^[-\U000F0000-\U000FFFFF]+\s*")


def _clean_line(raw):
    """Strip PDF icon/glyph characters from start of line, then strip whitespace."""
    return _ICON_RE.sub("", raw).strip()


def parse_name_and_title(header_lines):
    name, title = "", ""
    for raw in header_lines:
        line = _clean_line(raw)
        if not line:
            continue
        if _CONTACT_RE.search(line):
            continue
        if _ADDRESS_RE.match(line):
            continue
        # Skip very long lines (paragraphs, not a name/title)
        if len(line) > 80:
            continue
        if not name:
            name = line
        elif not title:
            title = line
            break
    return name, title


def parse_summary(lines):
    parts = []
    for l in lines:
        s = l.strip()
        if s and not _CONTACT_RE.search(s):
            parts.append(s)
    return " ".join(parts)


def parse_skills(lines):
    result = {
        "primary_expertise": "",
        "programming_languages": "",
        "methodologies": "",
        "other_skills": "",
    }
    leftovers = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        ll = stripped.lower()

        # Try to detect labelled lines like "Programming languages: Python, JS"
        colon_val = re.sub(r"^[^:]+:\s*", "", stripped) if ":" in stripped else ""

        if re.search(r"primary|main skill|core skill|area of expertise|speciali[sz]", ll):
            result["primary_expertise"] = colon_val or stripped
        elif re.search(r"programming|coding lang|scripting lang|dev language", ll):
            result["programming_languages"] = colon_val or stripped
        elif re.search(r"method|agile|scrum|devops|kanban|sdlc|framework", ll):
            result["methodologies"] = colon_val or stripped
        else:
            leftovers.append(stripped)

    if leftovers:
        result["other_skills"] = ", ".join(leftovers)

    return result


def parse_education(lines):
    entries = []
    cur = {"degree": "", "institute": "", "location": "", "year": ""}

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        ym = re.search(r"(\d{4})\s*[-–]\s*(\d{4}|present|current)", stripped, re.I)
        if ym:
            cur["year"] = ym.group(0)
            continue

        if re.search(r"bachelor|master|phd|doctor|associate|diploma|degree|b\.s|m\.s|b\.a|m\.a|mba|llb|bsc|msc", stripped, re.I):
            if any(cur.values()):
                entries.append(cur.copy())
                cur = {"degree": "", "institute": "", "location": "", "year": ""}
            cur["degree"] = stripped
        elif re.search(r"university|college|institute|school|academy|faculty|polytechnic", stripped, re.I):
            cur["institute"] = stripped
        elif re.search(r",\s*[A-Za-z]", stripped) and len(stripped) < 60:
            cur["location"] = stripped
        elif not cur["degree"]:
            cur["degree"] = stripped

    if any(cur.values()):
        entries.append(cur)

    return entries or [{"degree": "", "institute": "", "location": "", "year": ""}]


_KNOWN_LANGUAGES = {
    "english","bulgarian","german","french","spanish","italian","portuguese",
    "russian","polish","czech","slovak","romanian","hungarian","dutch","greek",
    "turkish","arabic","chinese","japanese","korean","hindi","swedish","danish",
    "norwegian","finnish","ukrainian","serbian","croatian","slovenian","macedonian",
    "albanian","bosnian","latvian","lithuanian","estonian","hebrew","persian","thai",
    "vietnamese","indonesian","malay",
}
_LEVEL_WORDS = {"native","fluent","proficient","intermediate","basic","elementary",
                "advanced","beginner","mother tongue","bilingual","professional",
                "working","conversational","limited","full","c2","c1","b2","b1","a2","a1"}


def parse_languages(lines):
    langs = []
    for line in lines:
        stripped = line.strip()
        if not stripped or _CONTACT_RE.search(stripped) or len(stripped) > 80:
            continue

        # Pattern 1: "Language: Level" or "Language - Level"
        m = re.match(r"^([A-Za-z\s\(\)]{2,35}?)\s*[:|-]\s*(.{2,60})$", stripped)
        if m:
            lang = m.group(1).strip()
            level = m.group(2).strip()
            if lang.lower() in _KNOWN_LANGUAGES:
                langs.append({"language": lang, "level": level})
                continue

        # Pattern 2: standalone known language name (with optional level word after)
        words = stripped.lower().split()
        if words and words[0] in _KNOWN_LANGUAGES:
            level = " ".join(w for w in words[1:] if w in _LEVEL_WORDS) or "Fluent"
            langs.append({"language": stripped.split()[0].capitalize(), "level": level})

    return langs or [{"language": "English", "level": "Professional working proficiency"}]


def parse_certifications(lines):
    certs = []
    for line in lines:
        stripped = re.sub(r"^[-•*►▸–]\s*", "", line.strip())
        if stripped and len(stripped) > 3:
            certs.append(stripped)
    return certs


# ---------------------------------------------------------------------------
# Work experience parser
# ---------------------------------------------------------------------------

_DATE_PAT = re.compile(
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|"
    r"january|february|march|april|june|july|august|september|october|november|december)"
    r"\.?\s*\d{4}\s*[-–—]\s*"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|"
    r"january|february|march|april|june|july|august|september|october|november|december"
    r"|present|current|now|till date|to date)\.?\s*\d{0,4}",
    re.I,
)
_DATE_PAT2 = re.compile(r"\d{4}\s*[-–—]\s*(?:\d{4}|present|current|now)", re.I)
# MM/YYYY – MM/YYYY  or  MM/YYYY – Present  (Denislav-style)
_DATE_PAT3 = re.compile(
    r"\d{1,2}[/\.]\d{4}\s*[-–—]\s*(?:\d{1,2}[/\.]\d{4}|present|current|now)",
    re.I,
)

_TITLE_KEYWORDS = re.compile(
    r"engineer|developer|manager|analyst|consultant|specialist|coordinator|"
    r"director|lead|architect|designer|admin|officer|head|executive|associate|"
    r"intern|technician|support|helpdesk|desktop|recruiter|hr |devops|"
    r"scientist|researcher|tester|qa |qa$|success|owner|scrum|agile|"
    r"sales|account|business|product|project|program|operations|marketing|"
    r"finance|controller|accountant|lawyer|attorney|advisor|strategist|"
    r"representative|agent|partner|senior|junior|principal|staff|chief|"
    r"vice president|vp |cto|ceo|coo|cfo|ciso",
    re.I,
)


def _find_date(line):
    # Try most-specific patterns first so MM/YYYY beats bare YYYY
    m = _DATE_PAT.search(line)   # Month-name YYYY – Month-name YYYY
    if m:
        return m.group(0).strip()
    m = _DATE_PAT3.search(line)  # MM/YYYY – MM/YYYY  (must come before _DATE_PAT2)
    if m:
        return m.group(0).strip()
    m = _DATE_PAT2.search(line)  # YYYY – YYYY  (bare year range)
    if m:
        return m.group(0).strip()
    return None


def parse_work_experience(lines):
    jobs = []
    cur = None
    prev_clean = ""   # track previous non-empty line for title-before-date pattern

    def flush():
        if cur and (cur["position"] or cur["responsibilities"]):
            jobs.append(cur.copy())

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        bullet = line[:2].strip() in ("-", "•", "*", "►", "▸", "–")
        clean = re.sub(r"^[-•*►▸–]\s*", "", line)

        date = _find_date(line)

        if date:
            flush()
            cur = {
                "project_name": "",
                "employer": "Confidential",
                "duration": date,
                "responsibilities": [],
                "position": "",
                "location": "",
            }
            # Title on same line (before date): e.g. "Software Engineer – Jan 2021 – Present"
            remainder = line.replace(date, "").strip(" -|–—,[]")
            if remainder and _TITLE_KEYWORDS.search(remainder) and len(remainder) < 80:
                cur["position"] = remainder
            # Title on the PREVIOUS line (most common layout): the line just before the date
            elif prev_clean and _TITLE_KEYWORDS.search(prev_clean) and len(prev_clean) < 80:
                cur["position"] = prev_clean
                # Remove that line from previous job's responsibilities if it was added
                if jobs and jobs[-1]["responsibilities"] and jobs[-1]["responsibilities"][-1] == prev_clean:
                    jobs[-1]["responsibilities"].pop()

        elif cur is None:
            pass  # Haven't found first job yet

        elif bullet:
            if clean:
                cur["responsibilities"].append(clean)

        elif not cur["position"] and _TITLE_KEYWORDS.search(clean) and len(clean) < 80:
            cur["position"] = clean

        elif re.search(r"(?i)(sofia|london|berlin|paris|amsterdam|bucharest|warsaw|prague|"
                        r"budapest|vienna|zurich|dubai|new york|singapore)", clean) and len(clean) < 60:
            cur["location"] = clean

        else:
            if clean and len(clean) > 10 and not _CONTACT_RE.search(clean):
                cur["responsibilities"].append(clean)

        if clean:
            prev_clean = clean

    flush()

    return jobs or [{
        "project_name": "",
        "employer": "Confidential",
        "duration": "",
        "responsibilities": [],
        "position": "",
        "location": "",
    }]


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def parse_cv(filepath):
    """
    Parse a CV file and return a structured dict.
    Raises RuntimeError with a user-friendly message if extraction fails.
    """
    raw = extract_text(filepath)
    secs = split_sections(raw)

    # Diagnostics: check if any real sections were found
    found_sections = [k for k in secs if k != "header" and secs[k]]

    name, job_title = parse_name_and_title(secs.get("header", []))

    # If we couldn't split into sections, everything lands in 'header'.
    # Fall back: treat the whole text as a flat blob and do best-effort extraction.
    if not found_sections:
        all_lines = [l.strip() for l in raw.splitlines() if l.strip()]
        name, job_title = parse_name_and_title(all_lines[:10])
        # Put remaining lines in summary as a starting point
        summary_text = " ".join(all_lines[2:15]) if len(all_lines) > 2 else ""
        return {
            "name":            name,
            "job_title":       job_title,
            "summary":         summary_text,
            "technologies":    {"primary_expertise": "", "programming_languages": "",
                                "methodologies": "", "other_skills": ""},
            "education":       [{"degree": "", "institute": "", "location": "", "year": ""}],
            "languages":       [{"language": "English", "level": "Professional working proficiency"}],
            "certifications":  [],
            "work_experience": [{"project_name": "", "employer": "Confidential", "duration": "",
                                 "responsibilities": [], "position": "", "location": ""}],
            "_warning": (
                "Could not detect CV sections automatically.\n"
                "The name and a summary snippet were pre-filled from the first lines.\n"
                "Please fill in the remaining fields manually."
            ),
        }

    jobs = parse_work_experience(secs.get("experience", []))

    # If no title was found in the header, try to infer it from the first job's position
    if not job_title and jobs and jobs[0].get("position"):
        job_title = jobs[0]["position"]

    return {
        "name":            name,
        "job_title":       job_title,
        "summary":         parse_summary(secs.get("summary", [])),
        "technologies":    parse_skills(secs.get("skills", [])),
        "education":       parse_education(secs.get("education", [])),
        "languages":       parse_languages(secs.get("languages", [])),
        "certifications":  parse_certifications(secs.get("certifications", [])),
        "work_experience": jobs,
    }
